import os
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


def doc_process(input_path, output_path):

    qsz_folder = os.path.join(input_path, '起诉状')
    wts_folder = os.path.join(input_path, '委托书')
    info_path = os.path.join(input_path, 'info.xlsx')

    qsz_output_shex = os.path.join(output_path, '起诉状_上海耳序')
    qsz_output_fjzy = os.path.join(output_path, '起诉状_福建智云')
    qsz_output_hnsx = os.path.join(output_path, '起诉状_海南申信')

    wts_output_shex = os.path.join(output_path, '委托书_上海耳序')
    wts_output_fjzy = os.path.join(output_path, '委托书_福建智云')
    wts_output_hnsx = os.path.join(output_path, '委托书_海南申信')

    phone_wanglei = '18916935832'
    phone_zhangliren = '13817213203'

    qsz_file_list = os.listdir(qsz_folder)
    wts_file_list = os.listdir(wts_folder)

    if not os.path.exists(os.path.join(output_path)):
        os.mkdir(os.path.join(output_path))

    if not os.path.exists(qsz_output_shex):
        os.mkdir(qsz_output_shex)
        
    if not os.path.exists(qsz_output_fjzy):
        os.mkdir(qsz_output_fjzy)
        
    if not os.path.exists(qsz_output_hnsx):
        os.mkdir(qsz_output_hnsx)

    if not os.path.exists(wts_output_shex):
        os.mkdir(wts_output_shex)
        
    if not os.path.exists(wts_output_fjzy):
        os.mkdir(wts_output_fjzy)
        
    if not os.path.exists(wts_output_hnsx):
        os.mkdir(wts_output_hnsx)

    df = pd.read_excel(info_path)

    num_shex = 0
    num_fjzy = 0
    num_hnsx = 0
    num_other_lawyer = 0

    # 改管辖法院，并存到对应融担公司的目录中
    for qsz_file in qsz_file_list:
        
        # 读取起诉状文件
        qsz_path = os.path.join(qsz_folder, qsz_file)
        
        # 加载Docx文件
        doc = Document(qsz_path)
        
        # 从文件名获取合同号
        contract_id = qsz_file.split('_')[1]
        
        # 从表格中找到对应合同号的管辖法院
        court = df[df['合同号']==contract_id]['管辖法院'].tolist()[0]

        for p in doc.paragraphs:
            if '人民法院' in p.text:

                # 将正确的管辖法院更新到Docx中（加run设置字体）
                p.text = ''
                run = p.add_run(court)
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

                # 调整字体大小
                font = p.style.font
                font.size = Pt(14)
        
        # 判断担保公司
        if '福建智云' in df[df['合同号']==contract_id]['融担公司'].tolist()[0]:
            output_folder = qsz_output_fjzy
            num_fjzy += 1
        elif '海南申信' in df[df['合同号']==contract_id]['融担公司'].tolist()[0]:
            output_folder = qsz_output_hnsx
            num_hnsx += 1
        else:
            output_folder = qsz_output_shex
            num_shex += 1
            
        # 保存文件到对应目录
        output_path = os.path.join(output_folder, qsz_file)
        doc.save(output_path)

    # 改律师名字和电话，并存到对应融担公司的目录中
    for wts_file in wts_file_list:
        
        # 读取委托书文件
        wts_path = os.path.join(wts_folder, wts_file)
        
        # 加载Docx文件
        doc = Document(wts_path)
        
        # 从文件名获取合同号
        contract_id = wts_file.split('_')[1]
        
        # 从表格中找到对应合同号的管辖法院
        lawyer = df[df['合同号']==contract_id]['承办律师'].tolist()[0]
        user = df[df['合同号']==contract_id]['用户姓名'].tolist()[0]

        # 替换律师和电话
        if lawyer != '王磊':
            num_other_lawyer += 1
            for p in doc.paragraphs:
                if '王磊' in p.text and user not in p.text:
                    text_new_name = p.text.replace('王磊', '张立人')
                    p.text = ''
                    # 加run用于修改字体
                    run = p.add_run(text_new_name)
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
                    # 调整字体
                    font = p.style.font
                    font.size = Pt(14)
                    
                if phone_wanglei in p.text:
                    text_new_phone = p.text.replace(phone_wanglei, phone_zhangliren)
                    p.text = ''
                    # 加run用于修改字体
                    run = p.add_run(text_new_phone)
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
                    # 调整字体
                    font = p.style.font
                    font.size = Pt(14)

                if '王磊' in p.text and user in p.text:
                    # 现委托 王磊 在我单位与 沈黎宾 追偿权纠纷案件中，作为我单位的委托代理人，代理权限如下：
                    p.text = ''
                    run = p.add_run('现委托')
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

                    run = p.add_run(' 张立人 ')
                    run.font.underline = True
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

                    run = p.add_run('在我单位与')
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

                    run = p.add_run(f' {user} ')
                    run.font.underline = True
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')

                    run = p.add_run('追偿权纠纷案件中，作为我单位的委托代理人，代理权限如下：')
                    run.font.name = "Arial"
                    run._element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')                

        # 判断担保公司
        if '福建智云' in df[df['合同号']==contract_id]['融担公司'].tolist()[0]:
            output_folder = wts_output_fjzy
        elif '海南申信' in df[df['合同号']==contract_id]['融担公司'].tolist()[0]:
            output_folder = wts_output_hnsx
        else:
            output_folder = wts_output_shex
            
        # 保存文件到对应目录
        output_path = os.path.join(output_folder, wts_file)
        doc.save(output_path)
    
    total_have = len(qsz_file_list)
    total_done = num_shex+num_fjzy+num_hnsx
    result = f'共 {total_have} 条\n完成 {total_done} 条\n\n上海耳序：共 {num_shex} 条 | 福建智云：共 {num_fjzy} 条 | 海南申信：共 {num_hnsx} 条\n\n \
王磊律师：共 {total_have-num_other_lawyer} 条 | 张立人律师：共 {num_other_lawyer} 条\n\n所有文档已完成自动编辑！'
    
    return result

if __name__ == '__main__':
    doc_process()